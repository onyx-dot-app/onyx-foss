"""Preview + compare harness for the build-session Markdown -> DOCX converter.

This is a developer tool for tuning ``onyx.server.features.build.session.md_to_docx``
toward the visual output that pandoc/pypandoc produces. It can:

* ``generate`` a ``.docx`` from a Markdown file using the in-tree mistune
  converter, and (when ``pypandoc`` is importable) the pandoc reference output
  alongside it, so the two can be opened side by side in Word.
* ``compare`` two ``.docx`` files structurally (paragraph-style histogram,
  per-paragraph style/text alignment, table and image counts) and print a
  similarity score, so progress while tuning is measurable without eyeballing.

Examples::

    # Generate <stem>.mistune.docx (+ <stem>.pypandoc.docx if pypandoc present)
    python -m scripts.build_docx_preview generate path/to/report.md

    # Structurally diff two documents
    python -m scripts.build_docx_preview compare a.docx b.docx

    # Do both: generate from markdown then compare mistune vs pypandoc
    python -m scripts.build_docx_preview all path/to/report.md

``pypandoc`` is intentionally not a project dependency (the whole point of the
converter is to avoid shipping the pandoc binary). Install it transiently when
tuning, e.g. ``uv run --with pypandoc-binary python -m scripts.build_docx_preview ...``.
"""

from __future__ import annotations

import argparse
import importlib
import sys
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

from onyx.server.features.build.session.md_to_docx import markdown_to_docx_bytes

_PYPANDOC_INSTALL_HINT = (
    "pypandoc not available; skipping the pandoc reference output. Re-run under "
    "`uv run --with pypandoc-binary ...` to generate it."
)


# --------------------------------------------------------------------------- #
# Generation
# --------------------------------------------------------------------------- #
def generate_mistune(md_text: str, out_path: Path) -> Path:
    out_path.write_bytes(markdown_to_docx_bytes(md_text))
    return out_path


def generate_pypandoc(md_text: str, out_path: Path) -> Path | None:
    """Render with pandoc for reference. Returns None if pypandoc is unavailable.

    pypandoc is intentionally not a project dependency, so it is imported
    dynamically and only when this dev tool is run under it.
    """
    try:
        pypandoc = importlib.import_module("pypandoc")
    except ImportError:
        print(_PYPANDOC_INSTALL_HINT, file=sys.stderr)
        return None
    pypandoc.convert_text(md_text, "docx", format="gfm", outputfile=str(out_path))
    return out_path


# --------------------------------------------------------------------------- #
# Structural inspection
# --------------------------------------------------------------------------- #
@dataclass
class DocStats:
    path: Path
    style_histogram: Counter[str]
    # (style, text) per non-empty paragraph, in document order.
    paragraphs: list[tuple[str, str]]
    num_tables: int
    num_embedded_images: int
    num_alt_text_images: int


def _count_embedded_images(document: DocxDocument) -> int:
    # Embedded pictures show up as <a:blip> / <pic:pic> inside the drawing tree.
    return len(document.element.findall(".//" + qn("pic:pic")))


def inspect(path: Path) -> DocStats:
    document = Document(str(path))
    paragraphs: list[tuple[str, str]] = []
    alt_text_images = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else "?"
        paragraphs.append((style, text))
        # The mistune converter renders images it cannot embed as "[image: ...]".
        if text.startswith("[image"):
            alt_text_images += 1
    return DocStats(
        path=path,
        style_histogram=Counter(style for style, _ in paragraphs),
        paragraphs=paragraphs,
        num_tables=len(document.tables),
        num_embedded_images=_count_embedded_images(document),
        num_alt_text_images=alt_text_images,
    )


# --------------------------------------------------------------------------- #
# Comparison
# --------------------------------------------------------------------------- #
def _paragraph_alignment_score(
    a: list[tuple[str, str]], b: list[tuple[str, str]]
) -> float:
    """Fraction of position-aligned paragraphs whose text AND style both match."""
    if not a and not b:
        return 1.0
    matches = sum(1 for pa, pb in zip(a, b) if pa == pb)
    return matches / max(len(a), len(b))


def _text_alignment_score(a: list[tuple[str, str]], b: list[tuple[str, str]]) -> float:
    """Fraction of position-aligned paragraphs whose text matches (ignoring style)."""
    if not a and not b:
        return 1.0
    matches = sum(1 for (_, ta), (_, tb) in zip(a, b) if ta == tb)
    return matches / max(len(a), len(b))


def compare(
    reference: DocStats, candidate: DocStats, *, max_mismatches: int = 25
) -> None:
    """Print a human-readable structural diff of candidate vs reference."""
    print(
        f"\n=== COMPARE ===\n  reference (target): {reference.path}\n  candidate (ours):   {candidate.path}\n"
    )

    print(
        "paragraphs (non-empty):"
        f" reference={len(reference.paragraphs)} candidate={len(candidate.paragraphs)}"
    )
    print(
        f"tables:           reference={reference.num_tables} candidate={candidate.num_tables}"
    )
    print(
        f"embedded images:  reference={reference.num_embedded_images} candidate={candidate.num_embedded_images}"
    )
    print(
        f"alt-text images:  reference={reference.num_alt_text_images} candidate={candidate.num_alt_text_images}"
    )

    print("\nparagraph-style histogram:")
    all_styles = sorted(set(reference.style_histogram) | set(candidate.style_histogram))
    width = max(len(s) for s in all_styles) if all_styles else 0
    for style in all_styles:
        r = reference.style_histogram.get(style, 0)
        c = candidate.style_histogram.get(style, 0)
        flag = "" if r == c else "  <-- differs"
        print(f"  {style:<{width}}  reference={r:<4} candidate={c:<4}{flag}")

    text_score = _text_alignment_score(reference.paragraphs, candidate.paragraphs)
    style_score = _paragraph_alignment_score(reference.paragraphs, candidate.paragraphs)
    print("\nsimilarity:")
    print(f"  text alignment (text only):     {text_score:6.1%}")
    print(f"  full alignment (text + style):  {style_score:6.1%}")

    print(
        f"\nfirst {max_mismatches} aligned paragraphs that differ (style and/or text):"
    )
    shown = 0
    for index, (ref_p, cand_p) in enumerate(
        zip(reference.paragraphs, candidate.paragraphs)
    ):
        if ref_p == cand_p:
            continue
        ref_style, ref_text = ref_p
        cand_style, cand_text = cand_p
        style_note = (
            "" if ref_style == cand_style else f" [{cand_style} != {ref_style}]"
        )
        print(f"  #{index}{style_note}")
        if ref_text != cand_text:
            print(f"     ref:  {ref_text[:90]}")
            print(f"     ours: {cand_text[:90]}")
        shown += 1
        if shown >= max_mismatches:
            break
    if shown == 0:
        print("  (none — paragraphs align exactly)")


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def _cmd_generate(md_path: Path, outdir: Path) -> tuple[Path, Path | None]:
    md_text = md_path.read_text(encoding="utf-8")
    outdir.mkdir(parents=True, exist_ok=True)
    stem = md_path.stem
    mistune_out = generate_mistune(md_text, outdir / f"{stem}.mistune.docx")
    print(f"wrote {mistune_out}")
    pypandoc_out = generate_pypandoc(md_text, outdir / f"{stem}.pypandoc.docx")
    if pypandoc_out:
        print(f"wrote {pypandoc_out}")
    return mistune_out, pypandoc_out


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    sub = parser.add_subparsers(dest="command", required=True)

    gen = sub.add_parser(
        "generate", help="markdown -> .docx (mistune, + pypandoc if available)"
    )
    gen.add_argument("markdown", type=Path)
    gen.add_argument("--outdir", type=Path, default=Path.cwd())

    cmp_ = sub.add_parser("compare", help="structurally diff two .docx files")
    cmp_.add_argument("reference", type=Path, help="the target output (e.g. pypandoc)")
    cmp_.add_argument("candidate", type=Path, help="our output (e.g. mistune)")

    all_ = sub.add_parser(
        "all", help="generate from markdown, then compare mistune vs pypandoc"
    )
    all_.add_argument("markdown", type=Path)
    all_.add_argument("--outdir", type=Path, default=Path.cwd())

    args = parser.parse_args(argv)

    if args.command == "generate":
        _cmd_generate(args.markdown, args.outdir)
        return 0

    if args.command == "compare":
        compare(inspect(args.reference), inspect(args.candidate))
        return 0

    if args.command == "all":
        mistune_out, pypandoc_out = _cmd_generate(args.markdown, args.outdir)
        if pypandoc_out is None:
            print(
                "cannot compare without the pypandoc reference output.", file=sys.stderr
            )
            return 1
        compare(inspect(pypandoc_out), inspect(mistune_out))
        return 0

    return 1


if __name__ == "__main__":
    raise SystemExit(main())
