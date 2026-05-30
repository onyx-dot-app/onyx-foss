"""Convert Markdown to a DOCX document using mistune + python-docx.

Used by the build session "export as DOCX" feature. ``mistune`` parses the
Markdown into an AST and ``python-docx`` writes the ``.docx``; both are
pure-Python, so the conversion needs no external binary.

Supported constructs (covering what the LLM-generated documents emit):
headings, bold/italic/strikethrough/inline-code, bulleted/numbered/nested
lists (with loose-list continuation paragraphs and preserved ordered-list
start values), blockquotes, fenced code blocks, GFM tables, hyperlinks
(carrying inherited inline formatting), images (rendered as alt text), inline
``<br>`` line breaks, HTML entities, and horizontal rules. Other raw HTML is
dropped rather than shown as literal markup.
"""

from dataclasses import dataclass
from dataclasses import replace
from html import unescape
from io import BytesIO
from typing import Any
from typing import cast

import mistune
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.opc.package import OpcPackage
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Twips
from docx.styles.style import ParagraphStyle
from docx.table import _Cell
from docx.text.paragraph import Paragraph

_MONOSPACE_FONT = "Courier New"
_CODE_FONT_SIZE = Pt(9)
# pandoc styles links with a "Hyperlink" character style: muted blue, no
# underline (unlike python-docx's brighter underlined default).
_STYLE_HYPERLINK = "Hyperlink"
_STYLE_ID_HYPERLINK = "Hyperlink"
_LINK_COLOR = RGBColor(0x4F, 0x81, 0xBD)
# python-docx ships built-in "List Bullet"/"List Number" styles plus numbered
# variants up to level 3 ("List Bullet 2", "List Bullet 3", ...). Deeper nesting
# reuses the level-3 style.
_MAX_LIST_LEVEL = 3
# pandoc indents each list level by 0.5" (720 twips) with a 0.25" hanging marker;
# python-docx's built-in list numbering indents only half as far.
_LIST_INDENT_PER_LEVEL = 720
_LIST_HANGING_INDENT = 360

# Paragraph styles, mirroring how pandoc's default reference.docx names and
# spaces its prose so the output reads like the previous pandoc export. The
# spacing/indent/heading values below are reproduced from that reference (plain
# measurements, not the file itself, which stays out of the repo for licensing).
_STYLE_BODY = "Body Text"
_STYLE_FIRST_PARAGRAPH = "First Paragraph"
_STYLE_COMPACT = "Compact"
_STYLE_IMAGE_CAPTION = "Image Caption"
_STYLE_BLOCK_TEXT = "Block Text"

_BODY_SPACE = Pt(9)  # Body Text: 180 twips before/after
_COMPACT_SPACE = Pt(1.8)  # Compact (tight lists): 36 twips
_BLOCK_TEXT_SPACE = Pt(5)  # Block Text (blockquote): 100 twips
_BLOCK_TEXT_INDENT = Inches(1 / 3)  # Block Text left/right: 480 twips
_HEADING_COLOR = RGBColor(0x0F, 0x47, 0x61)
_HEADING_SIZES = {1: Pt(20), 2: Pt(16), 3: Pt(14), 4: Pt(12), 5: Pt(11), 6: Pt(11)}
# Document default font/size, matching pandoc's reference (Aptos 12pt body,
# Aptos Display headings) instead of python-docx's Cambria 11pt default.
_BODY_FONT = "Aptos"
_HEADING_FONT = "Aptos Display"
_BODY_FONT_SIZE = Pt(12)
# pandoc emits no page margins, so Word renders its 1" default; python-docx's
# template uses 1.25" left/right. Set 1" all round to match the pandoc look.
_PAGE_MARGIN = Inches(1)

# Footnotes are written as a real Word footnotes part (python-docx has no native
# API for them), so [^n] citations become superscript references that Word links
# to page-bottom notes. Style ids are the spaceless form of the style names.
_STYLE_FOOTNOTE_TEXT = "Footnote Text"
_STYLE_FOOTNOTE_REFERENCE = "Footnote Reference"
_STYLE_ID_FOOTNOTE_TEXT = "FootnoteText"
_STYLE_ID_FOOTNOTE_REFERENCE = "FootnoteReference"
_FOOTNOTES_PARTNAME = "/word/footnotes.xml"
_FOOTNOTES_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
)
_FOOTNOTES_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Enable the GFM table/strikethrough/url plugins, plus footnotes, so the AST
# covers the Markdown features that appear in these documents.
_markdown_parser = mistune.create_markdown(
    renderer=None,
    plugins=["table", "strikethrough", "url", "footnotes"],
)

Node = dict[str, Any]


# Code points XML 1.0 forbids: C0 controls except tab/newline/CR, the UTF-16
# surrogate range, and U+FFFE/U+FFFF. Mapped to None for str.translate (C-speed).
_XML_INVALID_TRANSLATION = {
    **{code: None for code in range(0x20) if code not in (0x09, 0x0A, 0x0D)},
    **{code: None for code in range(0xD800, 0xE000)},
    0xFFFE: None,
    0xFFFF: None,
}


def _strip_invalid_xml_chars(text: str) -> str:
    """Drop characters XML 1.0 forbids (NULL and most C0 controls).

    LLM output occasionally contains them and python-docx raises when
    writing them, so they are stripped before parsing."""
    return text.translate(_XML_INVALID_TRANSLATION)


@dataclass(frozen=True)
class _Fmt:
    """Inline formatting flags carried down through nested inline nodes."""

    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False


def markdown_to_docx_bytes(md_text: str) -> bytes:
    """Render Markdown text to the bytes of a .docx file."""
    tokens = _markdown_parser(_strip_invalid_xml_chars(md_text))
    nodes: list[Node] = tokens if isinstance(tokens, list) else []

    document = Document()
    _apply_pandoc_styles(document)
    footnote_block = next(
        (node for node in nodes if node.get("type") == "footnotes"), None
    )
    footnotes = (
        _Footnotes(document, footnote_block.get("children", []))
        if footnote_block is not None
        else None
    )
    _render_blocks(document, nodes, footnotes)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _apply_pandoc_styles(document: DocxDocument) -> None:
    """Add/configure the prose styles, approximating pandoc's reference.docx.

    python-docx's bare default template puts everything in ``Normal``; pandoc
    instead distributes prose across ``Body Text``/``First Paragraph``, tight
    lists into ``Compact``, blockquotes into ``Block Text``, and image captions
    into ``Image Caption``. Defining the same styles here lets the renderer
    assign them so the document reads like the pandoc export.
    """
    for section in document.sections:
        section.left_margin = _PAGE_MARGIN
        section.right_margin = _PAGE_MARGIN
        section.top_margin = _PAGE_MARGIN
        section.bottom_margin = _PAGE_MARGIN

    styles = document.styles
    existing = {style.name for style in styles}

    normal = cast(ParagraphStyle, styles["Normal"])
    normal.font.name = _BODY_FONT
    normal.font.size = _BODY_FONT_SIZE
    # python-docx's template defaults to 1.15x line spacing; pandoc uses single.
    normal.paragraph_format.line_spacing = 1.0

    def ensure(name: str, base: str) -> ParagraphStyle:
        if name not in existing:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base]
            existing.add(name)
        return cast(ParagraphStyle, styles[name])

    body = cast(ParagraphStyle, styles[_STYLE_BODY])  # ships in the default template
    body.paragraph_format.space_before = _BODY_SPACE
    body.paragraph_format.space_after = _BODY_SPACE

    ensure(_STYLE_FIRST_PARAGRAPH, _STYLE_BODY)

    compact = ensure(_STYLE_COMPACT, _STYLE_BODY)
    compact.paragraph_format.space_before = _COMPACT_SPACE
    compact.paragraph_format.space_after = _COMPACT_SPACE

    block_text = ensure(_STYLE_BLOCK_TEXT, _STYLE_BODY)
    block_text.paragraph_format.space_before = _BLOCK_TEXT_SPACE
    block_text.paragraph_format.space_after = _BLOCK_TEXT_SPACE
    block_text.paragraph_format.left_indent = _BLOCK_TEXT_INDENT
    block_text.paragraph_format.right_indent = _BLOCK_TEXT_INDENT

    caption = ensure(_STYLE_IMAGE_CAPTION, "Caption")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ensure(_STYLE_FOOTNOTE_TEXT, "Normal")
    if _STYLE_FOOTNOTE_REFERENCE not in existing:
        reference = styles.add_style(_STYLE_FOOTNOTE_REFERENCE, WD_STYLE_TYPE.CHARACTER)
        reference.font.superscript = True

    if _STYLE_HYPERLINK not in existing:
        hyperlink = styles.add_style(_STYLE_HYPERLINK, WD_STYLE_TYPE.CHARACTER)
        hyperlink.font.color.rgb = _LINK_COLOR

    for level, size in _HEADING_SIZES.items():
        heading = styles[f"Heading {level}"]
        heading.font.name = _HEADING_FONT
        heading.font.size = size
        heading.font.color.rgb = _HEADING_COLOR
        # pandoc headings are coloured + sized, not bold; python-docx's are bold.
        heading.font.bold = False


# --------------------------------------------------------------------------- #
# Footnotes
# --------------------------------------------------------------------------- #
class _PartParent:
    """Minimal parent so a Paragraph built in the footnotes part resolves
    ``paragraph.part`` (used for hyperlink relationships) to that part."""

    def __init__(self, part: XmlPart) -> None:
        self.part = part


class _Footnotes:
    """Builds a Word footnotes part and wires up references to it.

    Word stores footnotes in a separate ``word/footnotes.xml`` part with two
    reserved entries (separator + continuation separator) followed by the real
    notes; the body points at one via ``<w:footnoteReference w:id=...>``.

    Word footnotes are 1:1 with their reference, but Markdown lets one note be
    cited from several places (``attrs["index"]`` repeats). So, like pandoc, a
    fresh Word footnote is emitted per *reference* (content duplicated for
    repeats), in reference order, each with a unique id.
    """

    def __init__(self, document: DocxDocument, definitions: list[Node]) -> None:
        self._definitions = {
            int(item.get("attrs", {}).get("index", 0)): item for item in definitions
        }
        self._next_id = 1
        self._element = parse_xml(f'<w:footnotes xmlns:w="{_W_NS}"/>')
        self._element.append(_separator_footnote(-1, "separator"))
        self._element.append(_separator_footnote(0, "continuationSeparator"))
        package: OpcPackage = document.part.package
        self._part = XmlPart(
            PackURI(_FOOTNOTES_PARTNAME),
            _FOOTNOTES_CONTENT_TYPE,
            self._element,
            package,
        )
        document.part.relate_to(self._part, _FOOTNOTES_REL_TYPE)
        self._parent = _PartParent(self._part)

    def add_reference(self, paragraph: Paragraph, index: int) -> None:
        """Insert a superscript reference and emit its footnote definition."""
        item = self._definitions.get(index)
        if item is None:
            return
        footnote_id = self._next_id
        self._next_id += 1
        paragraph._p.append(_reference_run("w:footnoteReference", footnote_id))
        self._element.append(self._build_footnote(footnote_id, item))

    def _build_footnote(self, footnote_id: int, item: Node) -> Any:
        """Render a footnote definition into a ``<w:footnote>`` element.

        Footnotes are almost always a single paragraph; each paragraph-like
        block becomes a footnote paragraph, with the reference mark + a space
        leading the first one.
        """
        footnote = OxmlElement("w:footnote")
        footnote.set(qn("w:id"), str(footnote_id))
        for position, block in enumerate(item.get("children", [])):
            p_element = OxmlElement("w:p")
            p_pr = OxmlElement("w:pPr")
            p_style = OxmlElement("w:pStyle")
            p_style.set(qn("w:val"), _STYLE_ID_FOOTNOTE_TEXT)
            p_pr.append(p_style)
            p_element.append(p_pr)
            paragraph = Paragraph(p_element, self._parent)
            if position == 0:
                p_element.append(_reference_run("w:footnoteRef", None))
                paragraph.add_run(" ")
            _add_runs(paragraph, block.get("children", []), _Fmt(), None)
            footnote.append(p_element)
        return footnote


def _separator_footnote(footnote_id: int, separator_tag: str) -> Any:
    footnote = OxmlElement("w:footnote")
    footnote.set(qn("w:type"), separator_tag)
    footnote.set(qn("w:id"), str(footnote_id))
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    run.append(OxmlElement(f"w:{separator_tag}"))
    paragraph.append(run)
    footnote.append(paragraph)
    return footnote


def _reference_run(mark_tag: str, footnote_id: int | None) -> Any:
    """A run carrying the footnote-reference character style and the mark.

    ``mark_tag`` is ``w:footnoteReference`` (body, needs an id) or
    ``w:footnoteRef`` (the mark inside the note itself).
    """
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), _STYLE_ID_FOOTNOTE_REFERENCE)
    run_props.append(style)
    run.append(run_props)
    mark = OxmlElement(mark_tag)
    if footnote_id is not None:
        mark.set(qn("w:id"), str(footnote_id))
    run.append(mark)
    return run


# --------------------------------------------------------------------------- #
# Block-level rendering
# --------------------------------------------------------------------------- #
def _set_paragraph_style(paragraph: Paragraph, style_name: str) -> None:
    """Set a paragraph's style by id, skipping python-docx's by-name lookup.

    ``add_paragraph(style=name)`` resolves the style by a linear, XML-parsing
    scan of every style and repeats it per paragraph, which dominates runtime on
    table/list-heavy documents. Built-in and added style ids are the spaceless
    form of the name, so set ``w:pStyle`` directly.
    """
    paragraph._p.get_or_add_pPr().get_or_add_pStyle().val = style_name.replace(" ", "")


def _add_styled_paragraph(document: DocxDocument, style_name: str) -> Paragraph:
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, style_name)
    return paragraph


def _render_blocks(
    document: DocxDocument, nodes: list[Node], footnotes: "_Footnotes | None"
) -> None:
    # Like pandoc: the first prose paragraph after any non-paragraph block (a
    # heading, list, table, blockquote, code, or the document start) uses "First
    # Paragraph"; consecutive prose paragraphs use "Body Text".
    first_para_pending = True
    for node in nodes:
        node_type = node.get("type")
        if node_type in ("blank_line", "newline"):
            continue
        if node_type == "footnotes":
            # Definitions live in the footnotes part (emitted per reference), not
            # the body.
            continue
        if node_type in ("paragraph", "block_text"):
            children = node.get("children", [])
            if _is_image_only(children):
                # An image renders as a caption; pandoc follows it with Body Text.
                _render_image_caption(document, children)
                first_para_pending = False
            else:
                style = _STYLE_FIRST_PARAGRAPH if first_para_pending else _STYLE_BODY
                paragraph = _add_styled_paragraph(document, style)
                _add_runs(paragraph, children, _Fmt(), footnotes)
                first_para_pending = False
            continue

        if node_type == "heading":
            level = min(int(node.get("attrs", {}).get("level", 1)), 6)
            paragraph = _add_styled_paragraph(document, f"Heading {level}")
            _add_runs(paragraph, node.get("children", []), _Fmt(), footnotes)
        elif node_type == "block_code":
            _render_code(document, node)
        elif node_type == "block_quote":
            _render_quote(document, node, footnotes)
        elif node_type == "list":
            _render_list(document, node, level=0, footnotes=footnotes)
        elif node_type == "thematic_break":
            _render_thematic_break(document)
        elif node_type == "table":
            _render_table(document, node, footnotes)
            # pandoc styles the paragraph after a table as Body Text, not First.
            first_para_pending = False
            continue
        elif "children" in node:
            # Unknown block wrapper: recurse so its content is not dropped.
            _render_blocks(document, node["children"], footnotes)
        first_para_pending = True


def _render_code(document: DocxDocument, node: Node) -> None:
    raw = str(node.get("raw", "")).rstrip("\n")
    paragraph = document.add_paragraph()
    for index, line in enumerate(raw.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = _MONOSPACE_FONT
        run.font.size = _CODE_FONT_SIZE


def _render_quote(
    document: DocxDocument, node: Node, footnotes: "_Footnotes | None"
) -> None:
    for child in node.get("children", []):
        if child.get("type") == "paragraph":
            paragraph = _add_styled_paragraph(document, _STYLE_BLOCK_TEXT)
            _add_runs(paragraph, child.get("children", []), _Fmt(), footnotes)
        else:
            _render_blocks(document, [child], footnotes)


def _is_image_only(children: list[Node]) -> bool:
    """True if the inline content is a single image (a standalone figure)."""
    meaningful = [
        child
        for child in children
        if child.get("type") not in ("softbreak", "linebreak")
        and not (child.get("type") == "text" and not str(child.get("raw", "")).strip())
    ]
    return len(meaningful) == 1 and meaningful[0].get("type") == "image"


def _render_image_caption(document: DocxDocument, children: list[Node]) -> None:
    """Render a standalone image as its alt text in the Image Caption style.

    Remote images are not fetched/embedded (pandoc does not either); the alt
    text is what reads in the document.
    """
    image = next(child for child in children if child.get("type") == "image")
    alt = _collect_text(image.get("children", []))
    paragraph = _add_styled_paragraph(document, _STYLE_IMAGE_CAPTION)
    paragraph.add_run(alt or "image")


def _render_list(
    document: DocxDocument,
    node: Node,
    level: int,
    footnotes: "_Footnotes | None",
) -> None:
    attrs = node.get("attrs", {})
    ordered = bool(attrs.get("ordered", False))
    # mistune exposes tight/loose as a top-level key on the list node.
    tight = bool(node.get("tight", True))
    style_level = min(level + 1, _MAX_LIST_LEVEL)
    list_style = "List Number" if ordered else "List Bullet"
    if style_level > 1:
        list_style = f"{list_style} {style_level}"
    continue_style = (
        "List Continue" if style_level == 1 else f"List Continue {style_level}"
    )
    start = int(attrs.get("start", 1))

    # Tight lists match pandoc's "Compact" style. Compact carries no list marker,
    # so numbering is applied directly from the built-in list style's definition
    # (which also supplies the indentation); each list gets its own instance so
    # ordered lists restart correctly. Loose lists keep the built-in list style.
    if tight:
        num_id = _create_list_numbering(document, list_style, start)
        item_style = _STYLE_COMPACT if num_id is not None else list_style
    else:
        # Ordered lists get their own numbering instance so each restarts at its
        # start value; bullets can rely on the built-in style.
        item_style = list_style
        num_id = (
            _create_list_numbering(document, list_style, start) if ordered else None
        )

    for item in node.get("children", []):
        if item.get("type") != "list_item":
            continue
        has_rendered_marker = False
        for child in item.get("children", []):
            child_type = child.get("type")
            if child_type in ("blank_line", "newline"):
                continue
            if child_type in ("block_text", "paragraph"):
                marker = not has_rendered_marker
                paragraph_style = item_style if marker else continue_style
                paragraph = _add_styled_paragraph(document, paragraph_style)
                # Indent 0.5" per level (matching pandoc); the marker line hangs.
                paragraph.paragraph_format.left_indent = Twips(
                    _LIST_INDENT_PER_LEVEL * (level + 1)
                )
                paragraph.paragraph_format.first_line_indent = Twips(
                    -_LIST_HANGING_INDENT if marker else 0
                )
                if marker and num_id is not None:
                    _apply_numbering(paragraph, num_id)
                _add_runs(paragraph, child.get("children", []), _Fmt(), footnotes)
                has_rendered_marker = True
            elif child_type == "list":
                _render_list(document, child, level + 1, footnotes)
            else:
                _render_blocks(document, [child], footnotes)


def _create_list_numbering(
    document: DocxDocument, list_style_name: str, start: int
) -> int | None:
    """Create a fresh numbering instance bound to a built-in list style.

    Returning a dedicated ``numId`` lets a paragraph in a non-list style (e.g.
    ``Compact``) still render list markers + indentation. A ``startOverride`` is
    always emitted (even for ``start == 1``): without it, multiple instances of
    the same abstract numbering continue one shared counter instead of each list
    restarting. Returns None if the style has no numbering definition (caller
    falls back to the list style).
    """
    style = document.styles[list_style_name]
    numbering = document.part.numbering_part.element
    abstract_num_id = _abstract_num_id_for_style(numbering, style.style_id)
    if abstract_num_id is None:
        return None

    num_ids = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) is not None
    ]
    next_num_id = max(num_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))

    abstract_num_id_el = OxmlElement("w:abstractNumId")
    abstract_num_id_el.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_id_el)

    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    lvl_override.append(start_override)
    num.append(lvl_override)

    numbering.append(num)
    return next_num_id


def _abstract_num_id_for_style(numbering: Any, style_id: str) -> str | None:
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract_num.findall(qn("w:lvl")):
            p_style = lvl.find(qn("w:pStyle"))
            if p_style is not None and p_style.get(qn("w:val")) == style_id:
                return abstract_num.get(qn("w:abstractNumId"))
    return None


def _apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.val = 0
    num_id_el = num_pr.get_or_add_numId()
    num_id_el.val = num_id


def _render_thematic_break(document: DocxDocument) -> None:
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    p_pr.append(borders)


def _render_table(
    document: DocxDocument, node: Node, footnotes: "_Footnotes | None"
) -> None:
    header_cells: list[Node] = []
    body_rows: list[list[Node]] = []
    for section in node.get("children", []):
        section_type = section.get("type")
        if section_type == "table_head":
            header_cells = section.get("children", [])
        elif section_type == "table_body":
            for row in section.get("children", []):
                body_rows.append(row.get("children", []))

    num_cols = len(header_cells) or (len(body_rows[0]) if body_rows else 0)
    if num_cols == 0:
        return

    # pandoc renders a borderless table (python-docx's default "Normal Table")
    # with only a rule under the header row, and small left/right cell padding.
    table = document.add_table(rows=0, cols=num_cols)
    _set_table_cell_margins(table)

    if header_cells:
        cells = table.add_row().cells
        for index, cell_node in enumerate(header_cells[:num_cols]):
            _fill_cell(cells[index], cell_node, bold=True, footnotes=footnotes)
        _underline_header_cells(cells)
    for row in body_rows:
        cells = table.add_row().cells
        for index, cell_node in enumerate(row[:num_cols]):
            _fill_cell(cells[index], cell_node, bold=False, footnotes=footnotes)

    _remove_fixed_cell_widths(table)


def _remove_fixed_cell_widths(table: Any) -> None:
    """Drop python-docx's fixed cell widths so columns auto-fit content.

    python-docx splits the full text width equally across columns; pandoc lets
    the table shrink to its content (leaving whitespace around small tables).
    """
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            if tc_pr is None:
                continue
            for tc_w in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(tc_w)


def _set_table_cell_margins(table: Any) -> None:
    """Apply pandoc's table cell padding (108 twips left/right, 0 top/bottom)."""
    margins = OxmlElement("w:tblCellMar")
    for edge, width in (("top", 0), ("left", 108), ("bottom", 0), ("right", 108)):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    table._tbl.tblPr.append(margins)


def _underline_header_cells(cells: Any) -> None:
    """Draw a single bottom rule under each header cell, like pandoc."""
    for cell in cells:
        borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "auto")
        borders.append(bottom)
        cell._tc.get_or_add_tcPr().append(borders)


_TABLE_CELL_ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}


def _fill_cell(
    cell: _Cell, cell_node: Node, bold: bool, footnotes: "_Footnotes | None"
) -> None:
    paragraph = cell.paragraphs[0]
    # pandoc uses the tight "Compact" style in cells and honours the column
    # alignment from the Markdown separator row (e.g. ``:--:`` -> centered).
    _set_paragraph_style(paragraph, _STYLE_COMPACT)
    alignment = _TABLE_CELL_ALIGN.get(str(cell_node.get("attrs", {}).get("align")))
    if alignment is not None:
        paragraph.alignment = alignment
    _add_runs(paragraph, cell_node.get("children", []), _Fmt(bold=bold), footnotes)


# --------------------------------------------------------------------------- #
# Inline rendering
# --------------------------------------------------------------------------- #
def _repair_paren_links(nodes: list[Node]) -> None:
    """Re-attach a balancing ``)`` that mistune split off a link destination.

    CommonMark allows balanced parentheses in a link/autolink destination (e.g.
    ``..._(novel)``), but mistune stops the URL at the first ``)`` and leaves it
    as following text. When a link's URL has unmatched ``(``, pull the matching
    ``)`` back from the next text node into the URL (and the visible text for an
    autolink), the way pandoc parses it.
    """
    for index, node in enumerate(nodes):
        if node.get("type") != "link" or index + 1 >= len(nodes):
            continue
        following = nodes[index + 1]
        if following.get("type") != "text":
            continue
        url = str(node.get("attrs", {}).get("url", ""))
        imbalance = url.count("(") - url.count(")")
        raw = str(following.get("raw", ""))
        take = 0
        while take < imbalance and take < len(raw) and raw[take] == ")":
            take += 1
        if take == 0:
            continue
        closing = ")" * take
        node.setdefault("attrs", {})["url"] = url + closing
        # For an autolink the visible text is the URL, and mistune split the same
        # ")" off it too. Detect that by the display ending with matching unclosed
        # "(" rather than comparing to the URL, which mistune may have re-encoded
        # (e.g. "\_" -> "%5C_"), so the closing reaches the display as well.
        children = node.get("children", [])
        if children and children[-1].get("type") == "text":
            display = str(children[-1].get("raw", ""))
            if display.count("(") - display.count(")") >= take:
                children[-1]["raw"] = display + closing
        following["raw"] = raw[take:]


def _add_runs(
    paragraph: Paragraph,
    nodes: list[Node],
    fmt: _Fmt,
    footnotes: "_Footnotes | None",
) -> None:
    _repair_paren_links(nodes)
    for node in nodes:
        node_type = node.get("type")
        if node_type == "text":
            # Decode HTML entities (e.g. ``&amp;``, ``&copy;``) the way a Markdown
            # renderer would; code spans below are intentionally left literal.
            _styled_run(paragraph, unescape(str(node.get("raw", ""))), fmt)
        elif node_type == "strong":
            _add_runs(
                paragraph, node.get("children", []), replace(fmt, bold=True), footnotes
            )
        elif node_type == "emphasis":
            _add_runs(
                paragraph,
                node.get("children", []),
                replace(fmt, italic=True),
                footnotes,
            )
        elif node_type == "strikethrough":
            _add_runs(
                paragraph,
                node.get("children", []),
                replace(fmt, strike=True),
                footnotes,
            )
        elif node_type == "codespan":
            _styled_run(paragraph, str(node.get("raw", "")), replace(fmt, code=True))
        elif node_type == "link":
            _add_hyperlink(paragraph, node, fmt, footnotes)
        elif node_type == "image":
            alt = _collect_text(node.get("children", []))
            _styled_run(paragraph, f"[image: {alt}]" if alt else "[image]", fmt)
        elif node_type == "footnote_ref":
            if footnotes is not None:
                footnotes.add_reference(
                    paragraph, int(node.get("attrs", {}).get("index", 0))
                )
        elif node_type == "softbreak":
            _styled_run(paragraph, " ", fmt)
        elif node_type == "linebreak":
            paragraph.add_run().add_break()
        elif node_type == "inline_html":
            _add_inline_html(paragraph, str(node.get("raw", "")))
        elif "children" in node:
            _add_runs(paragraph, node["children"], fmt, footnotes)
        elif "raw" in node:
            _styled_run(paragraph, unescape(str(node["raw"])), fmt)


def _styled_run(paragraph: Paragraph, text: str, fmt: _Fmt) -> None:
    run = paragraph.add_run(text)
    # Only set flags that are on, so plain runs don't emit redundant b="0"/i="0".
    if fmt.bold:
        run.bold = True
    if fmt.italic:
        run.italic = True
    if fmt.strike:
        run.font.strike = True
    if fmt.code:
        run.font.name = _MONOSPACE_FONT


def _add_hyperlink(
    paragraph: Paragraph, node: Node, fmt: _Fmt, footnotes: "_Footnotes | None"
) -> None:
    url = str(node.get("attrs", {}).get("url", ""))
    children = node.get("children", [])
    if not url:
        _add_runs(paragraph, children, fmt, footnotes)
        return

    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Render the label's inline content so nested bold/italic/code survive, then
    # move those runs into the hyperlink and tag each with the "Hyperlink"
    # character style (colour, no underline) like pandoc.
    start = len(paragraph._p)
    _add_runs(paragraph, children, fmt, footnotes)
    rendered = paragraph._p[start:]
    for element in rendered:
        paragraph._p.remove(element)
        if element.tag == qn("w:r"):
            _apply_hyperlink_style(element)
        hyperlink.append(element)
    if not rendered:
        # Empty label: fall back to the URL as the visible text.
        hyperlink.append(_hyperlink_text_run(url))
    paragraph._p.append(hyperlink)


def _apply_hyperlink_style(run: Any) -> None:
    run_props = run.find(qn("w:rPr"))
    if run_props is None:
        run_props = OxmlElement("w:rPr")
        run.insert(0, run_props)
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), _STYLE_ID_HYPERLINK)
    run_props.insert(0, run_style)  # rStyle is first in CT_RPr


def _hyperlink_text_run(text: str) -> Any:
    run = OxmlElement("w:r")
    _apply_hyperlink_style(run)
    text_el = OxmlElement("w:t")
    text_el.text = text
    if text != text.strip():
        text_el.set(qn("xml:space"), "preserve")
    run.append(text_el)
    return run


def _add_inline_html(paragraph: Paragraph, raw: str) -> None:
    normalized = raw.strip().lower()
    if normalized in ("<br>", "<br/>", "<br />"):
        paragraph.add_run().add_break()


def _collect_text(nodes: list[Node]) -> str:
    """Flatten inline nodes to plain text (for link labels and image alt text)."""
    parts: list[str] = []
    for node in nodes:
        if node.get("type") == "text":
            parts.append(unescape(str(node.get("raw", ""))))
        elif node.get("children"):
            parts.append(_collect_text(node["children"]))
        elif "raw" in node:
            parts.append(unescape(str(node["raw"])))
    return "".join(parts)
